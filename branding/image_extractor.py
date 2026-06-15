import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
}

def get_image_info(drawing_el, part) -> dict | None:
    """
    Extracts rId, width, and height from a w:drawing element.
    """
    try:
        # Find wp:extent to get size
        extent = drawing_el.find('.//wp:extent', NAMESPACES)
        width_emu = int(extent.get('cx')) if extent is not None else 0
        height_emu = int(extent.get('cy')) if extent is not None else 0
        
        # Convert EMUs to Pixels (approx 9525 EMUs per pixel at 96 DPI)
        width_px = int(width_emu / 9525) if width_emu else 0
        height_px = int(height_emu / 9525) if height_emu else 0
        
        # Find a:blip to get relation ID
        blip = drawing_el.find('.//a:blip', NAMESPACES)
        if blip is None:
            return None
            
        rId = blip.get(qn('r:embed'))
        if not rId:
            return None
            
        # Get image bytes
        image_part = part.related_parts[rId]
        image_bytes = image_part.blob
        
        # Generate unique hash/identifier for image content to detect repeats
        image_hash = hash(image_bytes)
        
        return {
            "rId": rId,
            "width": width_px,
            "height": height_px,
            "hash": image_hash,
            "bytes": image_bytes,
            "part": image_part
        }
    except Exception:
        return None

def extract_images_from_docx(doc) -> list[dict]:
    """
    Extracts all images from a DOCX document along with location metadata.
    """
    extracted_images = []
    image_counts = {} # hash -> count to calculate repeats
    
    # Helper to process a paragraph
    def process_p(p, location, page_approx):
        for run in p.runs:
            drawings = run.element.findall('.//w:drawing', NAMESPACES)
            for drawing in drawings:
                info = get_image_info(drawing, p.part)
                if info:
                    info["location"] = location
                    info["page"] = page_approx
                    info["paragraph"] = p
                    extracted_images.append(info)
                    image_counts[info["hash"]] = image_counts.get(info["hash"], 0) + 1

    # 1. Process Headers and Footers
    for section_idx, section in enumerate(doc.sections):
        # We estimate page based on section
        page_est = section_idx + 1
        
        # Helper to process a header or footer object
        def process_header_footer(hf_obj, loc_type):
            if not hf_obj:
                return
            for p in hf_obj.paragraphs:
                process_p(p, loc_type, page_est)
            for table in hf_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            process_p(p, loc_type, page_est)
                            
        process_header_footer(getattr(section, 'header', None), "header")
        process_header_footer(getattr(section, 'first_page_header', None), "header")
        process_header_footer(getattr(section, 'even_page_header', None), "header")
        
        process_header_footer(getattr(section, 'footer', None), "footer")
        process_header_footer(getattr(section, 'first_page_footer', None), "footer")
        process_header_footer(getattr(section, 'even_page_footer', None), "footer")

    # 2. Process Body Paragraphs
    # We estimate page number by distributing paragraphs (e.g. approx 30 paragraphs per page)
    for p_idx, p in enumerate(doc.paragraphs):
        page_est = (p_idx // 30) + 1
        # Treat first 5 paragraphs of every page as header location
        loc = "header" if (p_idx % 30 < 5) else "body"
        process_p(p, loc, page_est)

    # 3. Process Tables in Body
    for table_idx, table in enumerate(doc.tables):
        page_est = (table_idx // 3) + 1 # rough estimate
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    # Treat first table of every page as header location
                    loc = "header" if (table_idx % 3 == 0) else "body"
                    process_p(p, loc, page_est)
                    
    # 4. Attach frequency/repeat metadata to each image
    for img in extracted_images:
        img["repeat_count"] = image_counts.get(img["hash"], 1)
        
    return extracted_images
