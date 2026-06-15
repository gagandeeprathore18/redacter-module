import os
import sys
import docx
import cv2
import numpy as np

# Ensure project directories are in path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from branding.image_extractor import extract_images_from_docx
from branding.image_preprocessor import preprocess_image
from branding.ocr_detector import OCRDetector
from branding.branding_decision import BrandingDecisionEngine
from branding.logo_remover import remove_logo_inplace

def create_test_text_image(text: str, filepath: str):
    """
    Creates a temporary white image with custom text using OpenCV.
    """
    # Create white canvas
    img = np.ones((85, 220, 3), dtype=np.uint8) * 255
    # Write text
    cv2.putText(img, text, (10, 45), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (0, 0, 0), 1)
    # Save image
    cv2.imwrite(filepath, img)
    print(f"Created test image at: {filepath} with text: '{text}'")

def main():
    temp_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "temp_test_branding")
    os.makedirs(temp_dir, exist_ok=True)
    
    img_path = os.path.join(temp_dir, "test_logo.png")
    docx_path = os.path.join(temp_dir, "test_branding_input.docx")
    docx_output_path = os.path.join(temp_dir, "test_branding_output.docx")
    
    # 1. Create text-based logo
    create_test_text_image("Buckinghamshire", img_path)
    
    # 2. Create DOCX with image in Header
    doc = docx.Document()
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = "Document Header - "
    run = p.add_run()
    run.add_picture(img_path)
    
    # Save input document
    doc.save(docx_path)
    print(f"Created test DOCX at: {docx_path}")
    
    # 3. Reload and Process
    doc_proc = docx.Document(docx_path)
    extracted_images = extract_images_from_docx(doc_proc)
    
    assert len(extracted_images) == 1, f"Expected 1 image, found {len(extracted_images)}"
    img_meta = extracted_images[0]
    
    # Preprocess
    preprocessed_bytes = preprocess_image(img_meta["bytes"])
    
    # OCR Detector
    print("Running EasyOCR (this may take a few seconds on first run to download model files if any)...")
    ocr = OCRDetector()
    ocr_text = ocr.extract_text(preprocessed_bytes)
    print(f"Extracted OCR Text: '{ocr_text}'")
    
    # Branding Decision Engine
    decision_engine = BrandingDecisionEngine()
    should_remove, score, breakdown = decision_engine.evaluate_image(img_meta, ocr_text)
    
    print(f"Decision Results - Score: {score}, Should Remove: {should_remove}, Breakdown: {breakdown}")
    
    # Assertions
    assert "buckinghamshire" in ocr_text, "OCR failed to match 'buckinghamshire'"
    assert score >= 70, f"Expected score >= 70, got {score}"
    assert should_remove is True, "Expected image to be flagged for removal"
    
    # Remove
    remove_logo_inplace(img_meta)
    
    # Save processed document
    doc_proc.save(docx_output_path)
    print(f"Saved processed document to: {docx_output_path}")
    
    # 4. Verify output file has redacted image (1x1 pixel blob)
    doc_verify = docx.Document(docx_output_path)
    verify_images = extract_images_from_docx(doc_verify)
    assert len(verify_images) == 1, "Expected 1 image in output"
    
    # The output blob should be very small (1x1 PNG is about 67 bytes)
    blob_size = len(verify_images[0]["bytes"])
    print(f"Processed image blob size: {blob_size} bytes")
    assert blob_size < 200, f"Expected blob size < 200 bytes, got {blob_size}"
    
    print("\n--- BRANDING MODULE UNIT TESTS PASSED SUCCESSFULLY ---")

if __name__ == "__main__":
    main()
