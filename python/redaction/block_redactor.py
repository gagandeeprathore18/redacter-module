import docx

NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
}

def get_all_paragraph_runs(p):
    runs = []
    for r_el in p._element.findall('.//w:r', NAMESPACES):
        runs.append(docx.text.run.Run(r_el, p))
    return runs

def redact_cell(cell, is_docx=True):
    """
    Clear cell text to a single space (" ") in a formatting-preserving way.
    """
    if is_docx:
        for p in cell.paragraphs:
            runs = get_all_paragraph_runs(p)
            if not runs:
                p.text = " "
            else:
                runs[0].text = " "
                for run in runs[1:]:
                    run.text = ""
            # Clear text in nested tables recursively
            for table in cell.tables:
                for row in table.rows:
                    for c in row.cells:
                        redact_cell(c, is_docx=True)
    else:
        # PPTX cell
        if hasattr(cell, 'text_frame') and cell.text_frame:
            for p in cell.text_frame.paragraphs:
                if not p.runs:
                    p.text = " "
                else:
                    p.runs[0].text = " "
                    for run in p.runs[1:]:
                        run.text = ""
