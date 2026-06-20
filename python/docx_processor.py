import docx
import re
from PIL import Image
import io
import os
import json
from redact_engine import (
    redact_text, is_logo_match, redact_paragraph_runs,
    redact_paragraph_runs_with_pattern,
    TABLE_ROW_NAME_KEYWORD_PATTERN,
    SAFE_COLUMN_HEADER_PATTERN,
    TABLE_ROW_LOCATION_KEYWORD_PATTERN, SUBMISSION_LOCATION_PATTERN
)
from redaction.stop_patterns import should_stop_block

from redaction.table_parser import normalize_table
from redaction.target_detector import detect_targets
from redaction.relationship_detector import get_block_coordinates
from redaction.block_extractor import extract_block_cells
from redaction.block_redactor import redact_cell
from redaction.validator import validate_table_redaction
from redaction.docx_structure_analyzer import analyze_and_redact_runs
from redaction.header_footer_processor import process_docx_headers_footers
from redaction.protected_zone_detector import reset_protected_zone, set_protected_zone

NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
}

def get_all_paragraph_runs(p):
    """Ensure we capture all runs including those nested inside hyperlinks."""
    runs = list(p.runs)
    # Check for hyperlinks
    for child in p._element:
        if child.tag.endswith('hyperlink'):
            from docx.text.run import Run
            for run_el in child.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                runs.append(Run(run_el, p))
    return runs

def convert_image_to_png_bytes(image):
    buf = io.BytesIO()
    img = Image.open(io.BytesIO(image.blob))
    img.save(buf, format="PNG")
    return buf.getvalue()

def is_contact_or_address_block(text: str, is_header_footer: bool = False) -> tuple[bool, str]:
    if not text or not text.strip():
        return False, ""
        
    text_lower = text.lower()
    is_contact = any(k in text_lower for k in ["www", "http", "email", "@", "telephone", "tel", "fax", "contact"])
    from redact_engine import POSTAL_CODE_PATTERN
    has_postcode = bool(POSTAL_CODE_PATTERN.search(text))
    has_keyword = any(k in text_lower for k in ["street", "road", "avenue", "lane", "drive", "building", "house", "campus", "centre", "center", "park", "office", "gate", "london", "nottingham", "manchester", "oxford", "bucks", "leeds", "birmingham", "tel", "phone", "email", "@", "www", "http", "contact"])
    is_multiline = "\n" in text.strip()
    
    is_address = has_postcode and (has_keyword or is_multiline)
    has_explicit_office = any(k in text_lower for k in ["head office", "london office", "registered office", "office address"])
    
    if is_header_footer and is_contact:
        return True, "CONTACT_BLOCK"
        
    if is_address or has_explicit_office:
        return True, "ADDRESS_BLOCK"
        
    return False, ""

def redact_address_and_contact_paragraphs(paragraphs, is_header_footer=False):
    n = len(paragraphs)
    i = 0
    while i < n:
        matched = False
        for sz in range(min(5, n - i), 0, -1):
            window = paragraphs[i:i+sz]
            window_text = "\n".join(p.text for p in window).strip()
            
            should_remove, cls = is_contact_or_address_block(window_text, is_header_footer)
            if should_remove:
                for p in window:
                    for run in get_all_paragraph_runs(p):
                        run.text = " "
                try:
                    from redaction.redaction_audit import RedactionAudit
                    RedactionAudit.log({
                        "candidate": window_text,
                        "stage": "HEADER_REDACTION" if is_header_footer else "FINAL_DECISION",
                        "classification": cls,
                        "decision": "REDACT",
                        "page": None
                    })
                except Exception:
                    pass
                i += sz
                matched = True
                break
        if not matched:
            i += 1

def redact_and_remove_hyperlinks(p):
    """
    Finds w:hyperlink elements in the paragraph. If the target URL should be redacted,
    removes the hyperlink element from the paragraph completely.
    """
    from redact_engine import redact_text
    for child in list(p._element):
        if child.tag.endswith('hyperlink'):
            rId = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            if rId and rId in p.part.rels:
                rel = p.part.rels[rId]
                target = rel._target
                
                should_remove_link = False
                if "qualifi" in target.lower() or "qualify" in target.lower():
                    should_remove_link = True
                else:
                    from redact_engine import get_active_patterns, URL_OR_DOMAIN_PATTERN
                    for pattern in get_active_patterns():
                        if pattern.search(target):
                            should_remove_link = True
                            break
                    if not should_remove_link and URL_OR_DOMAIN_PATTERN.search(target):
                        for pattern in get_active_patterns():
                            if pattern.search(target):
                                should_remove_link = True
                                break
                
                if should_remove_link:
                    p._element.remove(child)

def process_paragraph(p, redact_all_names=False, blank_entire=False, context="", is_header_footer=False):
    """Process a paragraph's runs. If blank_entire is True, wipe all run text."""
    redact_and_remove_hyperlinks(p)
    runs = get_all_paragraph_runs(p)
    if not runs:
        return
    if blank_entire:
        for run in runs:
            run.text = " "
        return
    # Run structural run-combining analysis and redaction
    redact_paragraph_runs(runs, redact_all_names=redact_all_names, context=context, is_header_footer=is_header_footer)

def process_table(table, is_header_footer=False):
    # 1. Run Relationship-based Submission Location Redaction Engine
    grid = normalize_table(table)
    target_coords = detect_targets(grid)
    all_block_coords = set()
    for coord in target_coords:
        block = get_block_coordinates(grid, coord)
        all_block_coords.update(block)
        
    if all_block_coords:
        cells_to_redact = extract_block_cells(grid, all_block_coords)
        for cell in cells_to_redact:
            redact_cell(cell, is_docx=True)
        final_grid = normalize_table(table)
        validate_table_redaction(grid, final_grid, all_block_coords)

    # Get column headers if available
    headers = []
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            headers.append("".join(p.text for p in cell.paragraphs).strip())

    for row_idx, row in enumerate(table.rows):
        # Keyword scanning
        row_has_name_keyword = False
        row_has_location_keyword = False
        label_cell_idx = -1

        # The "label" cell is always the first unique cell in column 0.
        label_cell_text = ""
        if row.cells:
            label_cell_text = "".join(p.text for p in row.cells[0].paragraphs).strip()

        for idx, cell in enumerate(row.cells):
            # If this cell is part of the relationship-redacted block, skip scanning it
            if (row_idx, idx) in all_block_coords:
                continue
            cell_text = "".join(p.text for p in cell.paragraphs).strip()
            if TABLE_ROW_NAME_KEYWORD_PATTERN.search(cell_text):
                row_has_name_keyword = True
                label_cell_idx = idx
            if TABLE_ROW_LOCATION_KEYWORD_PATTERN.search(cell_text):
                row_has_location_keyword = True

        # Cell processing
        for idx, cell in enumerate(row.cells):
            # If cell was redacted by the relationship engine, skip it
            if (row_idx, idx) in all_block_coords:
                continue

            cell_text = "".join(p.text for p in cell.paragraphs).strip()
            set_protected_zone(cell_text)

            # Submission location rows: blank ALL cells entirely (label + value)
            if row_has_location_keyword:
                for p in cell.paragraphs:
                    set_protected_zone(p.text)
                    process_paragraph(p, blank_entire=True, is_header_footer=is_header_footer)
                for sub_table in cell.tables:
                    process_table(sub_table, is_header_footer=is_header_footer)
                continue

            # Run address and contact detection on cell paragraphs
            redact_address_and_contact_paragraphs(cell.paragraphs, is_header_footer=is_header_footer)

            # Only run redact_all_names if this is NOT the label cell
            redact_names_in_cell = row_has_name_keyword and (idx != label_cell_idx)
            for p in cell.paragraphs:
                set_protected_zone(p.text)
                process_paragraph(p, redact_all_names=redact_names_in_cell, context=label_cell_text if label_cell_text else "Name", is_header_footer=is_header_footer)

            # Recursively process sub-tables
            for sub_table in cell.tables:
                process_table(sub_table, is_header_footer=is_header_footer)

def should_stop_paragraph_block(text: str) -> bool:
    """Delegates to the shared stop-block checker."""
    return should_stop_block(text)

BLOCK_WIPE_LOCATION_PATTERN = re.compile(
    r'\b(?:'
    r'submission\s+(?:location|point|portal|platform|link|method|box|folder|area|mode|type|system|channel|url|address|site|page|form)'
    r'|submit(?:ted)?\s+(?:to|via|through|by)'
    r'|how\s+to\s+submit'
    r'|where\s+to\s+submit'
    r'|electronic(?:ally)?\s+submit(?:ted)?'
    r'|online\s+submission'
    r'|e-?submission'
    r'|upload\s+(?:location|link|portal|to)'
    r'|submission\s+details'
    r')\b',
    re.IGNORECASE
)

def is_paragraph_target(text: str) -> bool:
    if not text:
        return False
        
    from redaction.normalizer import normalize_text
    norm_text = normalize_text(text)
    
    from redaction.metadata_field_detector import is_metadata_field_to_keep
    if is_metadata_field_to_keep(text):
        return False
    
    # Do not target academic concepts or exclusions
    MUST_NEVER_BE_BUSINESS_FIELD = {
        "research",
        "research proposal",
        "research project",
        "research methods",
        "research skills",
        "methodology",
        "literature review",
        "findings",
        "analysis",
        "discussion",
        "recommendations",
        "presentation",
        "dissertation",
        "coursework",
        "assessment criteria",
        "learning outcomes",
        "recommended reading",
        "study hours",
        "guided study hours",
        "scheduled teaching hours",
        "assessment brief",
        "module guide",
        "reference list",
        "academic integrity",
        "confidentiality"
    }
    for excl in MUST_NEVER_BE_BUSINESS_FIELD:
        if excl in norm_text:
            return False

    if BLOCK_WIPE_LOCATION_PATTERN.search(text):
        return True
        
    from redaction.target_detector import BUSINESS_REMOVALS
    from redaction.metadata_field_detector import METADATA_LABELS
    from redaction.fuzzy_matcher import is_fuzzy_match
    for clean_target in list(BUSINESS_REMOVALS) + list(METADATA_LABELS):
        if is_fuzzy_match(norm_text, clean_target, threshold=85.0, partial=True):
            return True
    return False

def process_docx(input_path: str, output_path: str):
    os.environ["CURRENT_DOCUMENT_ID"] = os.path.basename(input_path)
    doc = docx.Document(input_path)
    
    # Register document context for the redaction debug logger
    from redaction.redaction_debug_logger import set_document_context
    set_document_context(document=os.path.basename(input_path), source="docx")

    # 0. Clear and Determine Issuing University first from logos
    from redaction.ownership_manager import clear_issuing_university, determine_issuing_university
    clear_issuing_university()
    
    images_to_remove = []
    image_ocr_results = {}
    try:
        from branding.image_extractor import extract_images_from_docx
        from branding.branding_decision import BrandingDecisionEngine
        from image_processing.pipeline import process_raster_image
        from redaction.redaction_audit import RedactionAudit
        
        extracted_images = extract_images_from_docx(doc)
        if extracted_images:
            decision_engine = BrandingDecisionEngine()
            
            for img in extracted_images:
                # Advanced Image Processing Pipeline
                res = process_raster_image(img["bytes"], img.get("location", ""))
                image_ocr_results[img["hash"]] = res
                ocr_text = res["ocr_text"]
                
                # Log pre-processing metrics
                RedactionAudit.log({
                    "candidate": f"[IMAGE rId={img['rId']}]",
                    "stage": "IMAGE_PRE_PROCESSING",
                    "classification": "IMAGE",
                    "decision": "KEEP",
                    "ocr_words_detected": res["ocr_words_detected"],
                    "char_map_entries": res["char_map_entries"],
                    "screenshot_ui_detected": res["screenshot_ui_detected"],
                    "screenshot_ui_cropped": 1 if res["screenshot_ui_detected"] else 0,
                })
                
                should_remove, score, breakdown = decision_engine.evaluate_image(img, ocr_text)
                print(f"Evaluated image {img['rId']}: score={score}, ocr='{ocr_text}', decision={'REMOVE' if should_remove else 'KEEP'}, breakdown={breakdown}")
                
                # Determine issuing university from logo OCR text
                determine_issuing_university(ocr_text)
                
                is_header_image = (img.get("location") == "header")
                if should_remove or is_header_image or is_logo_match(img["bytes"]):
                    images_to_remove.append(img)
    except Exception as ocr_err:
        print(f"Error during pre-scan: {ocr_err}")
    
    # 0.5. Scan pass for LLM Escalations
    from redaction.escalation_manager import clear_cache, register_candidate_scan, run_gpt_review, get_document_metrics
    from redaction.entity_classifier import classify_entity
    from redaction.ownership_manager import get_issuing_university
    
    clear_cache()
    reset_protected_zone()
    
    def scan_element_text(text, context="", in_table=False, is_header_footer=False):
        if not text or not text.strip():
            return
        if is_header_footer:
            from redaction.ownership_manager import scan_text_for_universities
            scan_text_for_universities(text)
        set_protected_zone(text)
        classification, action, reasons, score = classify_entity(text, context=context, in_table=in_table)
        register_candidate_scan(text, context, classification, action, score, reasons)

        # Scan for Proximity Names
        from redact_engine import NAME_PROXIMITY_PATTERN, NAME_PATTERN, TABLE_ROW_NAME_KEYWORD_PATTERN
        for match in NAME_PROXIMITY_PATTERN.finditer(text):
            name_str = match.group(1).strip()
            if name_str and len(name_str) > 2:
                classification, action, reasons, score = classify_entity(name_str, context=match.group(0), in_table=in_table)
                register_candidate_scan(name_str, match.group(0), classification, action, score, reasons)
                
        # Scan for general Names if name keyword exists
        if TABLE_ROW_NAME_KEYWORD_PATTERN.search(text):
            for match in NAME_PATTERN.finditer(text):
                name_str = match.group(0).strip()
                if name_str and len(name_str) > 2:
                    start_idx = max(0, match.start() - 120)
                    end_idx = min(len(text), match.end() + 120)
                    ctx = text[start_idx:end_idx]
                    classification, action, reasons, score = classify_entity(name_str, context=ctx, in_table=in_table)
                    register_candidate_scan(name_str, ctx, classification, action, score, reasons)

        # Scan for Date Candidates
        from redaction.date_time_detector import find_date_time_spans
        for start, end, matched_str, m_type in find_date_time_spans(text):
            if m_type == "DATE":
                matched_str = matched_str.strip()
                if matched_str and len(matched_str) > 2:
                    start_idx = max(0, start - 120)
                    end_idx = min(len(text), end + 120)
                    ctx = text[start_idx:end_idx]
                    classification, action, reasons, score = classify_entity(
                        matched_str, 
                        context=ctx, 
                        source_detector="DATE_CANDIDATE_PATTERN",
                        in_table=in_table
                    )
                    register_candidate_scan(matched_str, ctx, classification, action, score, reasons)

    # Paragraphs scan
    for p in doc.paragraphs:
        set_protected_zone(p.text)
        scan_element_text(p.text)

    # Tables scan
    def scan_table(t, is_header_footer=False):
        for row in t.rows:
            for cell in row.cells:
                cell_text = "".join(p.text for p in cell.paragraphs).strip()
                scan_element_text(cell_text, in_table=True, is_header_footer=is_header_footer)
                for p in cell.paragraphs:
                    scan_element_text(p.text, in_table=True, is_header_footer=is_header_footer)
                for sub_t in cell.tables:
                    scan_table(sub_t, is_header_footer=is_header_footer)

    for table in doc.tables:
        scan_table(table)
        
    # Headers & Footers scan
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                scan_element_text(p.text, is_header_footer=True)
            for table in section.header.tables:
                scan_table(table, is_header_footer=True)
        if section.footer:
            for p in section.footer.paragraphs:
                scan_element_text(p.text, is_header_footer=True)
            for table in section.footer.tables:
                scan_table(table, is_header_footer=True)
                
    # Run the GPT batch review if there are escalated candidates
    issuing_univ = get_issuing_university()
    doc_id = os.path.basename(input_path)
    run_gpt_review(issuing_university=issuing_univ, doc_id=doc_id)
    
    # Log passive telemetry metrics
    metrics = get_document_metrics(doc_id)
    print(f"Hybrid classification metrics: {json.dumps(metrics)}")
    
    # 1. Redact Paragraphs
    reset_protected_zone()
    
    # Run address and contact detection on all body paragraphs
    redact_address_and_contact_paragraphs(doc.paragraphs, is_header_footer=False)
    
    # Run address and contact detection on all header/footer paragraphs
    from redaction.header_footer_processor import get_docx_headers_footers
    for hf in get_docx_headers_footers(doc):
        redact_address_and_contact_paragraphs(hf.paragraphs, is_header_footer=True)
        
    in_location_block = False
    for p in doc.paragraphs:
        set_protected_zone(p.text)
        para_text = p.text.strip()
        if not para_text:
            if in_location_block:
                process_paragraph(p, blank_entire=True, is_header_footer=False)
            continue
            
        if is_paragraph_target(para_text):
            process_paragraph(p, blank_entire=True, is_header_footer=False)
            in_location_block = True
        elif in_location_block:
            if should_stop_paragraph_block(para_text):
                in_location_block = False
                process_paragraph(p, is_header_footer=False)
            else:
                process_paragraph(p, blank_entire=True, is_header_footer=False)
        else:
            process_paragraph(p, is_header_footer=False)
        
    # 2. Redact Tables
    for table in doc.tables:
        process_table(table, is_header_footer=False)
        
    # 3. Redact Headers and Footers using specialized processor with wrappers
    process_docx_headers_footers(
        doc,
        lambda p: process_paragraph(p, is_header_footer=True),
        lambda t: process_table(t, is_header_footer=True)
    )

    # 4. Redact Hyperlink URLs
    for rel_id, rel in doc.part.rels.items():
        if "hyperlink" in rel.reltype:
            rel._target = redact_text(rel._target)

    # 5. Remove marked logo images
    from branding.logo_remover import remove_logo_inplace
    for img in images_to_remove:
        try:
            remove_logo_inplace(img)

            # Phase 7 Log: Final Decision
            try:
                from redaction.redaction_audit import RedactionAudit
                classification = "HEADER_IMAGE" if img.get("location") == "header" else "UNIVERSITY_BRANDING"
                stage = "HEADER_REDACTION" if img.get("location") == "header" else "FINAL_DECISION"
                RedactionAudit.log({
                    "candidate": f"[IMAGE rId={img['rId']}]",
                    "stage": stage,
                    "classification": classification,
                    "decision": "REDACT",
                    "decision_source": "UNIVERSITY_BRANDING"
                })
            except Exception:
                pass
            # Phase 8 Log: Redaction Geometry
            try:
                from redaction.redaction_audit import RedactionAudit
                RedactionAudit.log({
                    "candidate": f"[IMAGE rId={img['rId']}]",
                    "page": None,
                    "stage": "REDACTION_GEOMETRY",
                    "bbox": None,
                    "bbox_width": None,
                    "bbox_height": None,
                    "ocr_text_inside_bbox": f"[IMAGE rId={img['rId']}]"
                })
            except Exception:
                pass
        except Exception as e:
            print(f"Error removing logo image: {e}")

    # 5.5. Partial / word-level redactions inside screenshots/images in DOCX
    from PIL import Image
    from PIL import ImageDraw
    import io
    if 'extracted_images' in locals() and extracted_images:
        for img in extracted_images:
            img_hash = img["hash"]
            if img in images_to_remove:
                continue
                
            if 'image_ocr_results' in locals() and img_hash in image_ocr_results:
                res = image_ocr_results[img_hash]
                char_mapper = res["char_mapper"]
                ocr_text = res["ocr_text"]
                if char_mapper and ocr_text.strip():
                    # Find all matches in ocr_text
                    from redact_engine import (
                        STUDENT_ID_PATTERN, EMAIL_PATTERN, PHONE_PATTERN, POSTAL_CODE_PATTERN,
                        NAME_PROXIMITY_PATTERN, is_likely_human_name
                    )
                    all_matches = []
                    patterns = [
                        (STUDENT_ID_PATTERN, "STUDENT_ID"),
                        (EMAIL_PATTERN, "EMAIL"),
                        (PHONE_PATTERN, "PHONE"),
                        (POSTAL_CODE_PATTERN, "POSTAL_CODE"),
                    ]
                    for pattern, classification in patterns:
                        for m in pattern.finditer(ocr_text):
                            all_matches.append((m.start(), m.end(), m.group(0), classification))
                            
                    for m in NAME_PROXIMITY_PATTERN.finditer(ocr_text):
                        name_str = m.group(1)
                        if is_likely_human_name(name_str, context=m.group(0)):
                            all_matches.append((m.start(1), m.end(1), name_str, "PERSON"))
                            
                    if all_matches:
                        try:
                            pil_img = Image.open(io.BytesIO(img["bytes"]))
                            draw = ImageDraw.Draw(pil_img)
                            modified = False
                            
                            for start, end, match_text, classification in all_matches:
                                sub_bbox = char_mapper.get_bbox_for_span(start, end)
                                if sub_bbox:
                                    # Apply adaptive padding
                                    from image_processing.redaction_padding import get_adaptive_padding
                                    padded_sub_bbox = get_adaptive_padding(sub_bbox, pil_img.size)
                                    
                                    # Sample background color
                                    from image_processing.background_sampler import sample_local_background
                                    local_bg = sample_local_background(pil_img, sub_bbox)
                                    
                                    px, py, pw, ph = padded_sub_bbox
                                    draw.rectangle([px, py, px + pw, py + ph], fill=local_bg)
                                    modified = True
                                    
                                    # Log the sub-redaction
                                    from redaction.redaction_audit import RedactionAudit
                                    RedactionAudit.log({
                                        "candidate": match_text,
                                        "stage": "FINAL_DECISION",
                                        "classification": classification,
                                        "decision": "REDACT",
                                        "background_sampled": True,
                                        "adaptive_padding": True
                                    })
                                    
                            if modified:
                                out_bytes_io = io.BytesIO()
                                fmt = pil_img.format if pil_img.format else "PNG"
                                pil_img.save(out_bytes_io, format=fmt)
                                img["part"].blob = out_bytes_io.getvalue()
                                print(f"Redacted sub-entities inside DOCX image rId={img['rId']} successfully")
                        except Exception as e:
                            print(f"Error redacting sub-entities inside DOCX image rId={img['rId']}: {e}")

    doc.save(output_path)

    # Write final audit summary report at the very end
    try:
        from redaction.redaction_audit import RedactionAudit
        from redaction.escalation_manager import LOGS_DIR
        summary = RedactionAudit.generate_summary(doc_id)
        summary_file = os.path.join(LOGS_DIR, "audit_summary.json")
        with open(summary_file, "w", encoding="utf-8") as sf:
            json.dump(summary, sf, indent=2)
    except Exception as re_err:
        print(f"Error generating final redaction audit report: {re_err}")

